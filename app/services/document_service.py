"""
Document processing service for extracting text from PDF and Word documents
"""
import logging
from typing import List, Optional, Tuple
from io import BytesIO
import pdfplumber
import PyPDF2

logger = logging.getLogger(__name__)

# Try to import python-docx for Word document processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Word document processing will be disabled.")


class DocumentService:
    """Service for processing PDF and Word documents"""
    
    @staticmethod
    def extract_text_from_pdf(pdf_bytes: bytes) -> str:
        """
        Extract text from PDF bytes
        
        Args:
            pdf_bytes: PDF file content as bytes
            
        Returns:
            Extracted text content
        """
        try:
            # Try pdfplumber first (better for tables and complex layouts)
            with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}, trying PyPDF2")
            
            try:
                # Fallback to PyPDF2
                pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
                text_parts = []
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                return "\n\n".join(text_parts)
            except Exception as e2:
                logger.error(f"PyPDF2 extraction also failed: {e2}")
                raise Exception(f"Failed to extract text from PDF: {e}, {e2}")
    
    @staticmethod
    def extract_text_from_docx(docx_bytes: bytes) -> str:
        """
        Extract text from Word document bytes
        
        Args:
            docx_bytes: Word document content as bytes
            
        Returns:
            Extracted text content
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
        
        try:
            doc = Document(BytesIO(docx_bytes))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to extract text from Word document: {e}")
            raise Exception(f"Failed to extract text from Word document: {e}")
    
    @staticmethod
    def extract_text_from_file(file_bytes: bytes, filename: str) -> Tuple[str, str]:
        """
        Extract text from file based on extension
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, file_type)
        """
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        if file_ext == 'pdf':
            text = DocumentService.extract_text_from_pdf(file_bytes)
            return text, 'pdf'
        elif file_ext in ['docx', 'doc']:
            text = DocumentService.extract_text_from_docx(file_bytes)
            return text, 'docx'
        else:
            raise ValueError(f"Unsupported file type: {file_ext}. Supported: pdf, docx, doc")
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Split text into chunks for processing
        
        Args:
            text: Text to chunk
            chunk_size: Target chunk size in characters
            overlap: Overlap between chunks in characters
            
        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings near the end
                for punct in ['. ', '.\n', '!\n', '?\n']:
                    last_punct = text.rfind(punct, start, end)
                    if last_punct != -1:
                        end = last_punct + len(punct)
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks

